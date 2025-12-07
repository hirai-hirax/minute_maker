from uuid import uuid4
from typing import List, Optional

from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field, ConfigDict
import os
import logging
import tempfile
import shutil
import json
from pathlib import Path
from io import BytesIO
from contextlib import contextmanager

from dotenv import load_dotenv
from openai import AzureOpenAI, OpenAI

import torch
import torchaudio
import numpy as np
import pandas as pd
from docx import Document as DocxDocument

# Monkeypatch huggingface_hub to support use_auth_token (removed in newer versions)
# This fixes the compatibility issue with speechbrain
import huggingface_hub
_original_hf_hub_download = huggingface_hub.hf_hub_download

def _patched_hf_hub_download(*args, **kwargs):
    if "use_auth_token" in kwargs:
        token = kwargs.pop("use_auth_token")
        if "token" not in kwargs:
            kwargs["token"] = token
    return _original_hf_hub_download(*args, **kwargs)

huggingface_hub.hf_hub_download = _patched_hf_hub_download

# Patch os.symlink for Windows to handle WinError 1314 (Privilege not held)
# SpeechBrain tries to symlink model files; fallback to copy if symlinks fail.
_original_symlink = os.symlink
def _patched_symlink(src, dst, *args, **kwargs):
    try:
        _original_symlink(src, dst, *args, **kwargs)
    except OSError as e:
        # Check for WinError 1314 (A required privilege is not held by the client)
        if getattr(e, 'winerror', None) == 1314:
            logger.warning(f"Symlink failed with WinError 1314. Falling back to copy for: {src} -> {dst}")
            if os.path.isdir(src):
                # Copy directory
                if os.path.exists(dst):
                     shutil.rmtree(dst)
                shutil.copytree(src, dst)
            else:
                # Copy file
                shutil.copy2(src, dst)
        else:
            raise

os.symlink = _patched_symlink

from speechbrain.pretrained import EncoderClassifier
from pydub import AudioSegment

load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Azure OpenAI Configuration
AZURE_OPENAI_ENDPOINT = os.environ.get("AZURE_OPENAI_ENDPOINT")
AZURE_OPENAI_API_KEY = os.environ.get("AZURE_OPENAI_API_KEY")
API_VERSION = "2025-03-01-preview"

# LLM Provider Configuration
LLM_PROVIDER = os.environ.get("LLM_PROVIDER", "azure")  # azure, ollama

# Ollama Configuration
OLLAMA_BASE_URL = os.environ.get("OLLAMA_BASE_URL", "http://localhost:11434/v1")
OLLAMA_MODEL = os.environ.get("OLLAMA_MODEL", "llama3.1")
OLLAMA_API_KEY = os.environ.get("OLLAMA_API_KEY", "ollama")

# OSS Whisper Configuration
WHISPER_PROVIDER = os.environ.get("WHISPER_PROVIDER", "azure")  # azure, faster-whisper
OSS_WHISPER_MODEL = os.environ.get("OSS_WHISPER_MODEL", "base")  # tiny, base, small, medium, large-v2, large-v3
OSS_WHISPER_DEVICE = os.environ.get("OSS_WHISPER_DEVICE", "cpu")  # cpu or cuda

# Directory Setup
BASE_DIR = Path(__file__).resolve().parent.parent
SPEAKERS_DIR = BASE_DIR / "data" / "speakers"
DATA_DIR = BASE_DIR / "data"
PROMPTS_FILE = DATA_DIR / "prompts.json"

SPEAKERS_DIR.mkdir(parents=True, exist_ok=True)
DATA_DIR.mkdir(parents=True, exist_ok=True)

# Default prompts (cannot be edited or deleted)
DEFAULT_PROMPTS = {
    "standard": {
        "id": "standard",
        "name": "標準校正",
        "description": "バランスの取れた標準的な要約・校正",
        "system_prompt": "あなたは議事録作成の専門家です。以下の会議の文字起こしテキストを読み、要約、決定事項、アクションアイテムを抽出してください。\n\n出力はJSON形式で、以下のキーを含めてください:\n- summary: 会議の要約\n- decisions: 決定事項のリスト\n- action_items: アクションアイテムのリスト",
        "is_default": True
    },
    "detailed": {
        "id": "detailed",
        "name": "詳細",
        "description": "詳細な分析と背景情報を含む要約",
        "system_prompt": "あなたは議事録作成の専門家です。以下の文字起こしを詳細に分析し、背景情報、議論の経緯を含めて要約してください。\n\n出力はJSON形式で、以下のキーを含めてください:\n- summary: 詳細な要約（背景含む）\n- decisions: 決定事項（経緯含む）\n- action_items: アクションアイテム（担当者明確化）",
        "is_default": True
    },
    "simple": {
        "id": "simple",
        "name": "簡潔",
        "description": "要点のみを箇条書きで",
        "system_prompt": "あなたは議事録作成の専門家です。以下の文字起こしから、要点のみを極めて簡潔に抽出してください。\n\n出力はJSON形式で、以下のキーを含めてください:\n- summary: 超簡潔な要約（3行以内）\n- decisions: 決定事項の箇条書き\n- action_items: アクションアイテムの箇条書き",
        "is_default": True
    }
}

# Global prompts data (default + custom)
PROMPTS_DATA = {}

def load_prompts():
    """Load prompts from JSON file, merging with defaults"""
    global PROMPTS_DATA
    PROMPTS_DATA = DEFAULT_PROMPTS.copy()
    
    if PROMPTS_FILE.exists():
        try:
            with open(PROMPTS_FILE, 'r', encoding='utf-8') as f:
                custom_prompts = json.load(f)
                PROMPTS_DATA.update(custom_prompts)
            logger.info(f"Loaded {len(custom_prompts)} custom prompts")
        except Exception as e:
            logger.error(f"Failed to load custom prompts: {e}")
    else:
        logger.info("No custom prompts file found, using defaults only")

def save_custom_prompts():
    """Save only custom prompts to JSON file"""
    custom_prompts = {k: v for k, v in PROMPTS_DATA.items() if not v.get("is_default", False)}
    try:
        with open(PROMPTS_FILE, 'w', encoding='utf-8') as f:
            json.dump(custom_prompts, f, ensure_ascii=False, indent=2)
        logger.info(f"Saved {len(custom_prompts)} custom prompts")
    except Exception as e:
        logger.error(f"Failed to save custom prompts: {e}")
        raise


# Chat model name - determined by provider
CHAT_MODEL_NAME = OLLAMA_MODEL if LLM_PROVIDER == "ollama" else "gpt-4o"


class MinuteBase(BaseModel):
    title: str = Field(..., description="Meeting or section title")
    summary: str = Field(..., description="Concise summary of the discussion")
    decisions: List[str] = Field(default_factory=list, description="Key decisions made")
    action_items: List[str] = Field(
        default_factory=list, description="Follow-up action items for attendees"
    )


class Minute(MinuteBase):
    id: str


app = FastAPI(title="Minute Maker API", version="0.1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("startup")
async def startup_event():
    """Pre-load models on startup to catch errors early"""
    try:
        logger.info("Loading prompts...")
        load_prompts()
        logger.info(f"Loaded {len(PROMPTS_DATA)} prompts")
        
        logger.info("Pre-loading speaker encoder on startup...")
        load_speaker_encoder()
        logger.info("Speaker encoder ready!")
        
        # Pre-load OSS Whisper model if configured
        if WHISPER_PROVIDER == "faster-whisper":
            logger.info(f"Pre-loading {WHISPER_PROVIDER} model ({OSS_WHISPER_MODEL})...")
            load_oss_whisper_model()
            logger.info(f"{WHISPER_PROVIDER} model ready!")
    except Exception as e:
        logger.error(f"Failed to pre-load models: {e}")
        logger.warning("Server will continue, but some features may not work")



_minutes: List[Minute] = [
    Minute(
        id=str(uuid4()),
        title="Project Kickoff",
        summary="Introduced project goals, timelines, and stakeholder expectations.",
        decisions=[
            "Adopt FastAPI for backend services",
            "Use React with Vite and TypeScript for the frontend",
        ],
        action_items=[
            "Set up repository scaffolding",
            "Draft initial UI for capturing meeting notes",
        ],
    ),
    Minute(
        id=str(uuid4()),
        title="API Contract Review",
        summary="Reviewed the REST contract for the minutes API and confirmed payload shapes.",
        decisions=["Expose /api/minutes for listing and creation"],
        action_items=["Add validation rules", "Document example payloads"],
    ),
]


@app.get("/", summary="Service health")
async def root() -> dict[str, str]:
    return {"message": "Minute Maker API is running"}


@app.get("/api/minutes", response_model=List[Minute], summary="List minutes")
async def list_minutes() -> List[Minute]:
    return _minutes


@app.get(
    "/api/minutes/{minute_id}",
    response_model=Minute,
    summary="Retrieve a single set of minutes",
)
async def get_minute(minute_id: str) -> Minute:
    for minute in _minutes:
        if minute.id == minute_id:
            return minute
    raise HTTPException(status_code=404, detail="Minute not found")


@app.post(
    "/api/minutes",
    response_model=Minute,
    status_code=201,
    summary="Create new meeting minutes",
)
async def create_minute(minute: MinuteBase) -> Minute:
    new_minute = Minute(id=str(uuid4()), **minute.model_dump())
    _minutes.append(new_minute)
    return new_minute


class TranscriptSegment(BaseModel):
    start: float
    end: float
    text: str
    speaker: str = ""

class ProcessingResult(BaseModel):
    id: str
    transcript: str
    segments: List[TranscriptSegment]
    summary: str
    speakers: List[str]


def _create_llm_client(api_version: str = API_VERSION):
    """Create LLM client based on configured provider (Azure OpenAI or Ollama)"""
    # Auto-detect provider if OLLAMA_BASE_URL is configured
    provider = LLM_PROVIDER
    if provider == "azure" and OLLAMA_BASE_URL and OLLAMA_BASE_URL != "http://localhost:11434/v1":
        # If Ollama URL is explicitly configured, use it
        provider = "ollama"
    
    if provider == "ollama":
        logger.info(f"Using Ollama LLM provider at {OLLAMA_BASE_URL} with model {OLLAMA_MODEL}")
        return OpenAI(
            base_url=OLLAMA_BASE_URL,
            api_key=OLLAMA_API_KEY,
        )
    else:
        # Default to Azure OpenAI
        if not AZURE_OPENAI_ENDPOINT or not AZURE_OPENAI_API_KEY:
            raise RuntimeError("AZURE_OPENAI_ENDPOINT and AZURE_OPENAI_API_KEY must be set")
        logger.info("Using Azure OpenAI LLM provider")
        return AzureOpenAI(
            azure_endpoint=AZURE_OPENAI_ENDPOINT,
            api_key=AZURE_OPENAI_API_KEY,
            api_version=api_version,
        )


# Backward compatibility alias
def _create_azure_client(api_version: str = API_VERSION) -> AzureOpenAI:
    """Deprecated: Use _create_llm_client instead"""
    return _create_llm_client(api_version)


# ---------------------------------------------------------------------------
# Speaker Identification Utilities
# ---------------------------------------------------------------------------

_SPEAKER_ENCODER = None

def load_speaker_encoder() -> EncoderClassifier:
    global _SPEAKER_ENCODER
    if _SPEAKER_ENCODER is None:
        logger.info("Loading SpeechBrain EncoderClassifier...")
        savedir = BASE_DIR / "tmp_model"
        
        # Check if model files exist locally (pre-downloaded)
        required_files = [
            "hyperparams.yaml",
            "embedding_model.ckpt",
            "mean_var_norm_emb.ckpt",
            "classifier.ckpt",
            "label_encoder.txt"
        ]
        
        files_exist = all((savedir / f).exists() for f in required_files)
        
        if files_exist:
            logger.info(f"Using pre-downloaded model from {savedir}")
            source = str(savedir)
        else:
            logger.info("Downloading model from Hugging Face...")
            logger.info("This may take a few minutes on first run (~100MB download)")
            source = "speechbrain/spkrec-ecapa-voxceleb"
            
            # Clear cache directory to ensure fresh download
            if savedir.exists():
                logger.info(f"Clearing existing model cache at {savedir}")
                try:
                    shutil.rmtree(savedir)
                except Exception as e:
                    logger.warning(f"Failed to clear cache: {e}")
            
            # Ensure directory exists
            savedir.mkdir(parents=True, exist_ok=True)
        
        try:
            _SPEAKER_ENCODER = EncoderClassifier.from_hparams(
                source=source,
                run_opts={"device": "cpu"},
                savedir=str(savedir)
            )
            logger.info("Speaker recognition model loaded successfully!")
        except Exception as e:
            logger.error(f"Failed to load speaker encoder: {e}", exc_info=True)
            raise HTTPException(
                status_code=500,
                detail=f"Failed to load speaker recognition model: {str(e)}"
            )
    
    return _SPEAKER_ENCODER


# ---------------------------------------------------------------------------
# OSS Whisper Model Loading
# ---------------------------------------------------------------------------

_WHISPER_MODEL = None

def load_oss_whisper_model():
    """Load OSS Whisper model (faster-whisper)"""
    global _WHISPER_MODEL
    if _WHISPER_MODEL is None:
        if WHISPER_PROVIDER == "faster-whisper":
            from faster_whisper import WhisperModel
            logger.info(f"Loading faster-whisper model: {OSS_WHISPER_MODEL}")
            compute_type = "int8" if OSS_WHISPER_DEVICE == "cpu" else "float16"
            _WHISPER_MODEL = WhisperModel(
                OSS_WHISPER_MODEL,
                device=OSS_WHISPER_DEVICE,
                compute_type=compute_type
            )
            logger.info(f"Loaded faster-whisper model on {OSS_WHISPER_DEVICE} with {compute_type}")
        else:
            raise ValueError(f"Unknown whisper provider: {WHISPER_PROVIDER}")
    return _WHISPER_MODEL


@contextmanager
def temp_file_path(data: bytes, suffix: str):
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(data)
        tmp_path = tmp.name
    try:
        yield tmp_path
    finally:
        if os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
            except Exception as e:
                logger.warning(f"Failed to delete temp file {tmp_path}: {e}")


def _compute_embedding_from_wav_bytes(wav_bytes: bytes) -> np.ndarray:
    with temp_file_path(wav_bytes, ".wav") as wav_path:
        # Load audio file
        waveform, sample_rate = torchaudio.load(wav_path)
    
    # Resample if necessary
    target_sr = 16000
    if sample_rate != target_sr:
        resampler = torchaudio.transforms.Resample(orig_freq=sample_rate, new_freq=target_sr)
        waveform = resampler(waveform)

    # Convert to single channel if necessary
    if waveform.shape[0] > 1:
        waveform = waveform.mean(dim=0, keepdim=True)
    
    encoder = load_speaker_encoder()
    waveform = waveform.to(dtype=torch.float32)
    
    # helper: classify_batch expects a batch, so we might need to ensure it's treated as one
    # 但是 encode_batch handles it. 
    with torch.no_grad():
        embedding = encoder.encode_batch(waveform)
    
    # embedding is (batch, 1, emb_dim) -> squeeze to (emb_dim,)
    return embedding.squeeze().cpu().numpy()


def _calculate_cosine_similarity(emb1: np.ndarray, emb2: np.ndarray) -> float:
    norm1 = np.linalg.norm(emb1)
    norm2 = np.linalg.norm(emb2)
    if norm1 == 0 or norm2 == 0:
        return 0.0
    return np.dot(emb1, emb2) / (norm1 * norm2)

def _apply_speaker_identification(
    audio_path: str,
    segments_data: List[dict],
    known_speakers: dict,
    threshold: float = 0.65
) -> (List[dict], set):
    """
    Returns updated segments (as dicts) and a set of all speakers found.
    """
    logger.info("Loading full audio for identification...")
    sound = AudioSegment.from_file(audio_path)
    
    updated_segments = []
    all_speakers = set()
    
    logger.info(f"Processing {len(segments_data)} segments...")

    for i, seg in enumerate(segments_data):
        start_ms = int(seg.get("start", 0) * 1000)
        end_ms = int(seg.get("end", 0) * 1000)
        
        # Avoid extremely short segments
        if end_ms - start_ms < 500:
            updated_segments.append(seg)
            if seg.get("speaker"): all_speakers.add(seg.get("speaker"))
            continue

        # Extract segment audio
        segment_audio = sound[start_ms:end_ms]
        seg_io = BytesIO()
        segment_audio.export(seg_io, format="wav")
        seg_io.seek(0)
        
        # Compute Embedding
        try:
            seg_emb = _compute_embedding_from_wav_bytes(seg_io.read())
            
            # Identify
            best_spk = seg.get("speaker", "")
            best_sim = -1.0
            best_spk_candidate = None
            
            for name, ref_emb in known_speakers.items():
                sim = _calculate_cosine_similarity(seg_emb, ref_emb)
                if sim > best_sim:
                    best_sim = sim
                    best_spk_candidate = name
            
            if best_spk_candidate and best_sim >= threshold:
                best_spk = best_spk_candidate
            
            seg["speaker"] = best_spk
            if best_spk:
                all_speakers.add(best_spk)
            
        except Exception as e:
            logger.warning(f"Error identification segment {i}: {e}")
        
        updated_segments.append(seg)
        
    return updated_segments, all_speakers


# ---------------------------------------------------------------------------
# Transcription Logic
# ---------------------------------------------------------------------------

def _transcribe_audio(audio_path: str, original_filename: str, model: str = "gpt-4o") -> List[dict]:
    """Transcribe audio using configured provider (Azure or OSS Whisper)"""
    if WHISPER_PROVIDER == "azure":
        return _transcribe_with_azure(audio_path, original_filename, model)
    elif WHISPER_PROVIDER == "faster-whisper":
        return _transcribe_with_oss_whisper(audio_path)
    else:
        raise ValueError(f"Unknown WHISPER_PROVIDER: {WHISPER_PROVIDER}")


def _transcribe_with_azure(audio_path: str, original_filename: str, model: str) -> List[dict]:
    """Transcribe using Azure OpenAI Whisper"""
    # Determine API version based on model
    api_version = API_VERSION
    if model == "whisper":
        api_version = "2024-06-01"
    
    client = _create_azure_client(api_version=api_version)
    logger.info(f"Starting Azure transcription with model: {model}, api_version: {api_version}")
    
    # Determine mime type/extension for the API
    ext = os.path.splitext(original_filename)[1]
    if not ext:
        ext = ".mp3"
    
    with open(audio_path, "rb") as audio_file:
        if model == "whisper":
            response = client.audio.transcriptions.create(
                model="whisper",
                file=(original_filename, audio_file, f"audio/{ext.lstrip('.')}"),
                response_format="verbose_json",
            )
            # Whisper returns valid JSON with segments but NO speaker info in verbose_json
            transcript_dict = response.model_dump()
            segments_data = transcript_dict.get("segments", [])
            
        else:
            # Default to gpt-4o-transcribe-diarize
            response = client.audio.transcriptions.create(
                model="gpt-4o-transcribe-diarize",
                file=(original_filename, audio_file, f"audio/{ext.lstrip('.')}"),
                response_format="diarized_json",
                chunking_strategy="auto",
            )
            transcript_dict = response.model_dump()
            segments_data = transcript_dict.get("segments", [])
    
    results = []
    for seg in segments_data:
        results.append({
            "start": seg.get("start", 0.0),
            "end": seg.get("end", 0.0),
            "text": seg.get("text", "").strip(),
            "speaker": seg.get("speaker", ""),  # Whisper will imply empty string here
        })
        
    return results


def _transcribe_with_oss_whisper(audio_path: str) -> List[dict]:
    """Transcribe using OSS faster-whisper"""
    logger.info(f"Starting OSS Whisper transcription with model: {OSS_WHISPER_MODEL}")
    whisper_model = load_oss_whisper_model()
    
    # Transcribe with faster-whisper
    segments, info = whisper_model.transcribe(
        audio_path,
        language="ja",  # Japanese - change to None for auto-detection
        beam_size=5,
        word_timestamps=False,
        vad_filter=True,  # Voice activity detection for better segmentation
    )
    
    logger.info(f"Detected language: {info.language} with probability {info.language_probability:.2f}")
    
    results = []
    for seg in segments:
        results.append({
            "start": seg.start,
            "end": seg.end,
            "text": seg.text.strip(),
            "speaker": ""  # No diarization in OSS Whisper - will be filled by SpeechBrain later
        })
    
    logger.info(f"Transcribed {len(results)} segments")
    return results


@app.post("/api/process_audio", response_model=ProcessingResult, summary="Process audio file")
async def process_audio(
    file: UploadFile = File(...),
    model: str = Form("gpt-4o")  # 'gpt-4o' or 'whisper'
):
    # Generate ID (for response tracking, not file storage)
    process_id = str(uuid4())
    suffix = Path(file.filename).suffix
    
    # Create temporary file for processing
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
            shutil.copyfileobj(file.file, tmp_file)
            tmp_path = tmp_file.name
        
        logger.info(f"Processing audio file: {file.filename} (process_id: {process_id})")
        
        # Run transcription with diarization logic
        segments_data = _transcribe_audio(tmp_path, file.filename, model)
        
        # Load system speakers
        sys_speakers = list(SPEAKERS_DIR.glob("*.npy"))
        known_speakers = {}
        if sys_speakers:
            logger.info(f"Loading {len(sys_speakers)} system speakers")
            for spk_file in sys_speakers:
                try:
                    known_speakers[spk_file.stem] = np.load(spk_file)
                except Exception as e:
                    logger.warning(f"Failed to load speaker {spk_file}: {e}")
            
            # Apply identification if we have known speakers
            if known_speakers:
                segments_data, speakers = _apply_speaker_identification(tmp_path, segments_data, known_speakers)
        
        # Convert to Pydantic models
        segments = [TranscriptSegment(**seg) for seg in segments_data]
        
        # Construct full transcript text
        full_transcript = ""
        speakers = set()
        for seg in segments:
            speaker_label = f"[{seg.speaker}] " if seg.speaker else ""
            full_transcript += f"[{seg.start:.2f} - {seg.end:.2f}] {speaker_label}{seg.text}\n"
            if seg.speaker:
                speakers.add(seg.speaker)
        
        return ProcessingResult(
            id=process_id,
            transcript=full_transcript.strip(),
            segments=segments,
            summary="Summary generation not yet implemented.",
            speakers=sorted(list(speakers))
        )
    except Exception as e:
        logger.error(f"Transcription failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        # Always clean up temporary file
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
                logger.info(f"Cleaned up temporary file: {tmp_path}")
            except Exception as e:
                logger.warning(f"Failed to delete temp file {tmp_path}: {e}")



@app.post("/api/create_speaker_embedding", summary="Create speaker embedding .npy file")
async def create_speaker_embedding(file: UploadFile = File(...)):
    content = await file.read()
    try:
        logger.info(f"Creating speaker embedding for file: {file.filename}")
        
        # Convert to WAV using pydub to handle various formats (mp3, m4a, etc.)
        # Load from memory
        audio = AudioSegment.from_file(BytesIO(content))
        
        # Export to WAV memory buffer
        wav_io = BytesIO()
        audio.export(wav_io, format="wav")
        wav_bytes = wav_io.getvalue()
        
        # Load encoder (may trigger download on first call)
        logger.info("Loading speaker encoder for embedding computation...")
        load_speaker_encoder()  # This will use the cached global instance
        
        embedding = _compute_embedding_from_wav_bytes(wav_bytes)
        logger.info(f"Successfully computed embedding (shape: {embedding.shape})")
        
        bio = BytesIO()
        np.save(bio, embedding)
        bio.seek(0)
        
        filename = Path(file.filename).stem + ".npy"
        
        return StreamingResponse(
            bio, 
            media_type="application/octet-stream",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Embedding creation failed: {e}", exc_info=True)
        raise HTTPException(
            status_code=500, 
            detail=f"Failed to create speaker embedding: {str(e)}"
        )


class SpeakerRegistration(BaseModel):
    start: float
    end: float


@app.post("/api/register_speaker", summary="Register speaker embedding from audio segment")
async def register_speaker(
    audio: UploadFile = File(...),
    start: float = Form(...),
    end: float = Form(...),
    speaker_name: str = Form(...)
):
    """Register a speaker by extracting embedding from a time segment of the uploaded audio file.
    
    Args:
        audio: Audio file to process
        start: Start time in seconds
        end: End time in seconds
        speaker_name: Name to save the speaker embedding as
    """
    tmp_path = None
    try:
        # Save uploaded file to temporary location
        suffix = Path(audio.filename).suffix
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
            shutil.copyfileobj(audio.file, tmp_file)
            tmp_path = tmp_file.name
        
        # Load audio segment
        start_ms = int(start * 1000)
        end_ms = int(end * 1000)
        
        # Verify duration
        if end_ms - start_ms < 500:
            raise HTTPException(status_code=400, detail="Segment too short for embedding (min 0.5s)")

        sound = AudioSegment.from_file(tmp_path)
        segment_audio = sound[start_ms:end_ms]
        
        seg_io = BytesIO()
        segment_audio.export(seg_io, format="wav")
        seg_io.seek(0)
        
        # Compute embedding
        embedding = _compute_embedding_from_wav_bytes(seg_io.read())
        
        # Save embedding
        safe_name = "".join([c for c in speaker_name if c.isalnum() or c in (' ', '_', '-')]).strip()
        if not safe_name:
            safe_name = "unknown_speaker"
            
        save_path = SPEAKERS_DIR / f"{safe_name}.npy"
        np.save(save_path, embedding)
        
        logger.info(f"Speaker '{safe_name}' registered successfully from audio segment")
        return {"message": f"Speaker {safe_name} registered successfully", "path": str(save_path)}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Registration failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        # Clean up temporary file
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
                logger.info(f"Cleaned up temporary file: {tmp_path}")
            except Exception as e:
                logger.warning(f"Failed to delete temp file {tmp_path}: {e}")


@app.get("/api/speakers", summary="List all registered speakers")
async def list_speakers():
    speakers = []
    for spk_file in SPEAKERS_DIR.glob("*.npy"):
        speakers.append({"name": spk_file.stem, "path": str(spk_file)})
    return sorted(speakers, key=lambda x: x["name"])


@app.post("/api/speakers", summary="Add a new speaker from audio or embedding file")
async def add_speaker(
    file: UploadFile = File(...),
    name: str = Form(...)
):
    try:
        content = await file.read()
        
        # Determine file type based on extension
        file_ext = Path(file.filename).suffix.lower()
        
        if file_ext == '.npy':
            # Direct .npy embedding file upload
            logger.info(f"Processing .npy embedding file: {file.filename}")
            
            # Load and validate the embedding
            bio = BytesIO(content)
            embedding = np.load(bio)
            
            # Validate embedding shape (should be 1D array)
            if embedding.ndim != 1:
                raise HTTPException(
                    status_code=400, 
                    detail=f"Invalid embedding shape: expected 1D array, got shape {embedding.shape}"
                )
            
            logger.info(f"Loaded embedding with shape: {embedding.shape}")
            
        else:
            # Audio file upload - compute embedding
            logger.info(f"Processing audio file: {file.filename}")
            
            # Convert to WAV using pydub to handle various formats (mp3, m4a, etc.)
            audio = AudioSegment.from_file(BytesIO(content))
            
            # Export to WAV memory buffer
            wav_io = BytesIO()
            audio.export(wav_io, format="wav")
            wav_bytes = wav_io.getvalue()
            
            embedding = _compute_embedding_from_wav_bytes(wav_bytes)
        
        # Save embedding with safe name
        safe_name = "".join([c for c in name if c.isalnum() or c in (' ', '_', '-')]).strip()
        if not safe_name:
            raise HTTPException(status_code=400, detail="Invalid speaker name")
            
        save_path = SPEAKERS_DIR / f"{safe_name}.npy"
        np.save(save_path, embedding)
        
        logger.info(f"Speaker '{safe_name}' saved successfully to {save_path}")
        return {"message": f"Speaker {safe_name} added successfully", "name": safe_name}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to add speaker: {e}", exc_info=True)
        # Return more detailed error
        raise HTTPException(status_code=500, detail=f"Processing failed: {str(e)}")


@app.delete("/api/speakers/{name}", summary="Delete a registered speaker")
async def delete_speaker(name: str):
    # Security check: basic strict filename
    safe_name = "".join([c for c in name if c.isalnum() or c in (' ', '_', '-')]).strip()
    if safe_name != name:
         raise HTTPException(status_code=400, detail="Invalid speaker name format")

    target_path = SPEAKERS_DIR / f"{safe_name}.npy"
    if not target_path.exists():
        raise HTTPException(status_code=404, detail="Speaker not found")
        
    try:
        os.unlink(target_path)
        return {"message": f"Speaker {name} deleted"}
    except Exception as e:
        logger.error(f"Failed to delete speaker: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/identify_speakers", response_model=ProcessingResult, summary="Identify speakers using embeddings")
async def identify_speakers(
    audio: UploadFile = File(...),
    transcript_json: str = Form(..., description="JSON string of transcript segments"),
    embedding_files: List[UploadFile] = File(default=[]),
    threshold: float = Form(0.65, description="Similarity threshold for identification"),
):
    """Identify speakers in audio segments using registered speaker embeddings.
    
    Args:
        audio: Audio file to analyze (required)
        transcript_json: JSON string containing transcript segments with timestamps
        embedding_files: Optional additional speaker embedding files (.npy)
        threshold: Similarity threshold for speaker identification (default: 0.65)
    """
    # Parse transcript
    try:
        segments_data = json.loads(transcript_json)
        # Validate minimal structure
        if not isinstance(segments_data, list):
            raise ValueError("Transcript must be a list of segments")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Invalid transcript JSON: {e}")

    # Load Embeddings
    known_speakers = {}
    
    # 1. Load system speakers
    sys_speakers = list(SPEAKERS_DIR.glob("*.npy"))
    logger.info(f"Loading {len(sys_speakers)} system speakers from {SPEAKERS_DIR}")
    for spk_file in sys_speakers:
        try:
            emb = np.load(spk_file)
            known_speakers[spk_file.stem] = emb
        except Exception as e:
            logger.warning(f"Failed to load system speaker {spk_file}: {e}")

    # 2. Load provided files
    for ef in embedding_files:
        try:
            content = await ef.read()
            bio = BytesIO(content)
            # Assuming strictly .npy files
            emb = np.load(bio)
            # Use filename (without extension) as speaker name
            name = Path(ef.filename).stem
            # Clean up name similar to reference code
            for delimiter in ['‗', '_']:
                if delimiter in name:
                    name = name.split(delimiter)[0]
                    break
            known_speakers[name] = emb
        except Exception as e:
            logger.warning(f"Failed to load embedding {ef.filename}: {e}")

    if not known_speakers:
        # If no valid embeddings, just return original
        # Reconstruct full transcript
        full_transcript = ""
        segments = []
        speakers = set()
        for seg in segments_data:
            s = TranscriptSegment(**seg)
            segments.append(s)
            speaker_label = f"[{s.speaker}] " if s.speaker else ""
            full_transcript += f"[{s.start:.2f} - {s.end:.2f}] {speaker_label}{s.text}\n"
            if s.speaker:
                speakers.add(s.speaker)
        
        return ProcessingResult(
            id=str(uuid4()),
            transcript=full_transcript.strip(),
            segments=segments,
            summary="No valid embeddings provided.",
            speakers=sorted(list(speakers))
        )

    # Process audio file using temporary storage
    tmp_path = None
    try:
        suffix = Path(audio.filename).suffix
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
            shutil.copyfileobj(audio.file, tmp_file)
            tmp_path = tmp_file.name
        
        logger.info(f"Identifying speakers in audio file: {audio.filename}")
        updated_segments, all_speakers = _apply_speaker_identification(
            tmp_path, segments_data, known_speakers, threshold
        )
        
        # Reconstruct Response
        segments_models = [TranscriptSegment(**s) for s in updated_segments]
        full_transcript = ""
        for s in segments_models:
            speaker_label = f"[{s.speaker}] " if s.speaker else ""
            full_transcript += f"[{s.start:.2f} - {s.end:.2f}] {speaker_label}{s.text}\n"
        
        return ProcessingResult(
            id=str(uuid4()),
            transcript=full_transcript.strip(),
            segments=segments_models,
            summary="Speaker identification applied.",
            speakers=sorted(list(all_speakers))
        )
        
    except Exception as e:
        logger.error(f"Speaker identification failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        # Always clean up temporary file
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
                logger.info(f"Cleaned up temporary file: {tmp_path}")
            except Exception as e:
                logger.warning(f"Failed to delete temp file {tmp_path}: {e}")


class SummarizeRequest(BaseModel):
    transcript: str
    prompt_id: str = "standard"
    llm_provider: Optional[str] = None  # Override provider: 'azure' or 'ollama'
    ollama_base_url: Optional[str] = None  # Ollama URL
    ollama_model: Optional[str] = None  # Ollama model name
    azure_model: Optional[str] = None  # Azure OpenAI model/deployment name

class SummarizeResponse(BaseModel):
    # 基本フィールド(後方互換性のため保持)
    summary: Optional[str] = None
    action_items: Optional[List[str]] = None
    decisions: Optional[List[str]] = None
    
    # 任意のカスタムフィールドを許可
    model_config = ConfigDict(extra='allow')

@app.get("/api/prompts", summary="Get available prompt presets")
async def get_prompts():
    return [
        {
            "id": k, 
            "name": v["name"], 
            "description": v["description"],
            "system_prompt": v["system_prompt"],
            "is_default": v.get("is_default", False)
        }
        for k, v in PROMPTS_DATA.items()
    ]

class PromptCreate(BaseModel):
    name: str = Field(..., min_length=1, max_length=100)
    description: str = Field(..., min_length=1, max_length=200)
    system_prompt: str = Field(..., min_length=1)

class PromptUpdate(BaseModel):
    name: Optional[str] = Field(None, min_length=1, max_length=100)
    description: Optional[str] = Field(None, min_length=1, max_length=200)
    system_prompt: Optional[str] = Field(None, min_length=1)

@app.post("/api/prompts", summary="Create a new custom prompt")
async def create_prompt(prompt: PromptCreate):
    # Generate unique ID
    prompt_id = f"custom_{str(uuid4())[:8]}"
    
    # Ensure uniqueness
    while prompt_id in PROMPTS_DATA:
        prompt_id = f"custom_{str(uuid4())[:8]}"
    
    # Add to data
    PROMPTS_DATA[prompt_id] = {
        "id": prompt_id,
        "name": prompt.name,
        "description": prompt.description,
        "system_prompt": prompt.system_prompt,
        "is_default": False
    }
    
    # Save to file
    try:
        save_custom_prompts()
        return {"message": "Prompt created successfully", "id": prompt_id}
    except Exception as e:
        # Rollback
        del PROMPTS_DATA[prompt_id]
        raise HTTPException(status_code=500, detail=f"Failed to save prompt: {str(e)}")

@app.put("/api/prompts/{prompt_id}", summary="Update a custom prompt")
async def update_prompt(prompt_id: str, prompt: PromptUpdate):
    # Check if prompt exists
    if prompt_id not in PROMPTS_DATA:
        raise HTTPException(status_code=404, detail="Prompt not found")
    
    # Check if it's a default prompt
    if PROMPTS_DATA[prompt_id].get("is_default", False):
        raise HTTPException(status_code=403, detail="Cannot edit default prompts")
    
    # Update fields
    if prompt.name is not None:
        PROMPTS_DATA[prompt_id]["name"] = prompt.name
    if prompt.description is not None:
        PROMPTS_DATA[prompt_id]["description"] = prompt.description
    if prompt.system_prompt is not None:
        PROMPTS_DATA[prompt_id]["system_prompt"] = prompt.system_prompt
    
    # Save to file
    try:
        save_custom_prompts()
        return {"message": "Prompt updated successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save prompt: {str(e)}")

@app.delete("/api/prompts/{prompt_id}", summary="Delete a custom prompt")
async def delete_prompt(prompt_id: str):
    # Check if prompt exists
    if prompt_id not in PROMPTS_DATA:
        raise HTTPException(status_code=404, detail="Prompt not found")
    
    # Check if it's a default prompt
    if PROMPTS_DATA[prompt_id].get("is_default", False):
        raise HTTPException(status_code=403, detail="Cannot delete default prompts")
    
    # Delete from data
    del PROMPTS_DATA[prompt_id]
    
    # Save to file
    try:
        save_custom_prompts()
        return {"message": "Prompt deleted successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save prompts: {str(e)}")


@app.post("/api/generate_summary", response_model=SummarizeResponse, summary="Generate summary from transcript")
async def generate_summary(
    transcript: str = Form(...),
    prompt_id: str = Form("standard"),
    llm_provider: Optional[str] = Form(None),
    ollama_base_url: Optional[str] = Form(None),
    ollama_model: Optional[str] = Form(None),
    azure_model: Optional[str] = Form(None),
    reference_files: List[UploadFile] = File(default=[])
):
    """
    Generate summary from transcript with optional reference materials.
    
    Args:
        transcript: The meeting transcript text
        prompt_id: ID of the prompt preset to use
        llm_provider: LLM provider ('azure' or 'ollama')
        ollama_base_url: Base URL for Ollama (if using Ollama)
        ollama_model: Model name for Ollama (if using Ollama)
        azure_model: Azure OpenAI model/deployment name (if overriding default)
        reference_files: Optional reference documents (.docx, .xlsx, .pptx, .pdf, .txt)
    """
    from .document_parser import extract_text_from_file
    
    # Create client with overridden settings if provided
    if llm_provider == "ollama" and ollama_base_url and ollama_model:
        logger.info(f"Using Ollama for summarization: {ollama_base_url} with model {ollama_model}")
        client = OpenAI(
            base_url=ollama_base_url,
            api_key="ollama"
        )
        model_name = ollama_model
    else:
        client = _create_llm_client()
        model_name = azure_model or CHAT_MODEL_NAME
        if azure_model:
            logger.info(f"Using Azure model override for summarization: {azure_model}")
    
    prompt_config = PROMPTS_DATA.get(prompt_id, PROMPTS_DATA["standard"])
    system_prompt = prompt_config["system_prompt"]
    
    # Process reference files
    reference_texts = []
    if reference_files:
        logger.info(f"Processing {len(reference_files)} reference files")
        for ref_file in reference_files:
            try:
                file_bytes = await ref_file.read()
                extracted_text = extract_text_from_file(ref_file.filename, file_bytes)
                reference_texts.append(f"--- Reference: {ref_file.filename} ---\n{extracted_text}")
                logger.info(f"Successfully extracted text from {ref_file.filename}")
            except Exception as e:
                logger.warning(f"Failed to extract text from {ref_file.filename}: {e}")
                # Continue with other files instead of failing completely
    
    # Build user prompt with transcript and optional reference materials
    user_prompt = f"""
    Transcript:
    {transcript}
    """
    
    if reference_texts:
        user_prompt += f"""
    
    Additional Reference Materials:
    {chr(10).join(reference_texts)}
    
    Please use the above reference materials as additional context when generating the summary.
    """
    
    try:
        completion = client.chat.completions.create(
            model=model_name,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            response_format={"type": "json_object"}
        )
        
        content = completion.choices[0].message.content

        # Normalize all fields from LLM response
        llm_data = json.loads(content)
        
        # Helper function to normalize list fields
        def normalize_list_field(field_value):
            """Normalize various list formats to List[str]"""
            result = []
            if isinstance(field_value, list):
                for item in field_value:
                    if isinstance(item, dict):
                        # Extract meaningful value from dict
                        if len(item) == 1:
                            result.append(str(list(item.values())[0]))
                        else:
                            result.append(" ".join(str(v) for v in item.values()))
                    elif isinstance(item, str):
                        result.append(item)
                    else:
                        result.append(str(item))
            elif isinstance(field_value, str):
                # Split by newlines if multi-line
                if "\n" in field_value:
                    result = [line.strip() for line in field_value.split("\n") if line.strip()]
                else:
                    result = [field_value]
            return result
        
        # Prepare response data
        response_data = {}
        
        # Process standard fields
        if "summary" in llm_data:
            response_data["summary"] = llm_data.pop("summary")
        
        if "action_items" in llm_data:
            response_data["action_items"] = normalize_list_field(llm_data.pop("action_items"))
        
        if "decisions" in llm_data:
            response_data["decisions"] = normalize_list_field(llm_data.pop("decisions"))
        
        # Add all remaining fields as custom fields (dynamic)
        for key, value in llm_data.items():
            if isinstance(value, list):
                response_data[key] = normalize_list_field(value)
            elif isinstance(value, str):
                response_data[key] = value
            elif isinstance(value, dict):
                # Convert dict to string representation
                response_data[key] = json.dumps(value, ensure_ascii=False)
            else:
                response_data[key] = str(value)
        
        return SummarizeResponse(**response_data)
    except Exception as e:
        logger.error(f"Summarization failed: {e}")
        # Fallback if JSON parsing fails or model fails
        return SummarizeResponse(
            summary="Failed to generate summary.",
            action_items=[],
            decisions=[]
        )


# ---------------------------------------------------------------------------
# File Download (Word/Excel)
# ---------------------------------------------------------------------------

class DownloadRequest(BaseModel):
    title: str
    summary: Optional[str] = None
    decisions: Optional[List[str]] = None
    action_items: Optional[List[str]] = None
    segments: List[TranscriptSegment]
    
    # \u30ab\u30b9\u30bf\u30e0\u30d5\u30a3\u30fc\u30eb\u30c9\u3092\u8a31\u53ef
    model_config = ConfigDict(extra='allow')


def _generate_docx(data: DownloadRequest) -> BytesIO:
    """Generate a Word document from minute data (mojiokoshi7.py準拠)"""
    from datetime import datetime
    
    doc = DocxDocument()
    
    # タイトル
    doc.add_heading('議事録', 0)
    doc.add_heading(data.title, level=1)
    
    # 作成日時
    creation_time = datetime.now().strftime('%Y年%m月%d日 %H:%M')
    doc.add_paragraph(f'作成日時: {creation_time}')
    doc.add_paragraph('')  # 空行
    
    # 統計情報
    doc.add_heading('統計情報', level=1)
    stats_lines = [
        f'総セグメント数: {len(data.segments)}件',
        f'決定事項: {len(data.decisions)}件',
        f'アクションアイテム: {len(data.action_items)}件',
    ]
    
    # 話者数をカウント
    speakers = set(seg.speaker for seg in data.segments if seg.speaker)
    if speakers:
        stats_lines.append(f'話者数: {len(speakers)}名')
    
    # 総時間を計算
    if data.segments:
        total_duration = max(seg.end for seg in data.segments)
        minutes = int(total_duration // 60)
        seconds = int(total_duration % 60)
        stats_lines.append(f'総時間: {minutes}分{seconds}秒')
    
    for stat in stats_lines:
        doc.add_paragraph(stat, style='List Bullet')
    doc.add_paragraph('')  # 空行
    
    # 要約セクション
    if data.summary:
        doc.add_heading('概要', level=1)
        doc.add_paragraph(data.summary)
        doc.add_paragraph('')  # 空行
    
    # 決定事項セクション
    if data.decisions:
        doc.add_heading('決定事項', level=1)
        for i, decision in enumerate(data.decisions, 1):
            doc.add_paragraph(f'{i}. {decision}', style='List Number')
        doc.add_paragraph('')  # 空行
    
    # アクションアイテムセクション
    if data.action_items:
        doc.add_heading('アクションアイテム', level=1)
        for i, action in enumerate(data.action_items, 1):
            doc.add_paragraph(f'{i}. {action}', style='List Number')
        doc.add_paragraph('')  # 空行
    
    # 文字起こし詳細セクション（mojiokoshi7.py形式）
    if data.segments:
        doc.add_heading('文字起こし', level=1)
        
        # 話者ごとにグループ化（mojiokoshi7.pyと同じロジック）
        merged_segments = []
        current_speaker = None
        current_texts = []
        current_start = None
        current_end = None
        
        for seg in data.segments:
            speaker = seg.speaker if seg.speaker else "不明"
            
            if speaker != current_speaker:
                if current_texts:
                    merged_segments.append({
                        "speaker": current_speaker,
                        "text": " ".join(current_texts),
                        "start": current_start,
                        "end": current_end
                    })
                current_speaker = speaker
                current_texts = [seg.text]
                current_start = seg.start
                current_end = seg.end
            else:
                current_texts.append(seg.text)
                current_end = seg.end  # 継続的に終了時刻を更新
        
        # 最後のグループを追加
        if current_texts:
            merged_segments.append({
                "speaker": current_speaker,
                "text": " ".join(current_texts),
                "start": current_start,
                "end": current_end
            })
        
        # 形式: [開始時刻-終了時刻]（話者）テキスト内容
        for merged in merged_segments:
            speaker = merged["speaker"]
            text = merged["text"]
            start_time = f"{int(merged['start']//60):02d}:{int(merged['start']%60):02d}"
            end_time = f"{int(merged['end']//60):02d}:{int(merged['end']%60):02d}"
            doc.add_paragraph(f"[{start_time}-{end_time}]（{speaker}）{text}")
    
    # カスタムフィールド(動的)
    # Pydantic extra='allow'で追加されたフィールドを処理
    custom_fields = {}
    if hasattr(data, '__pydantic_extra__'):
        custom_fields = data.__pydantic_extra__ or {}
    
    # 標準フィールド以外を表示
    standard_fields = {'title', 'summary', 'decisions', 'action_items', 'segments'}
    for field_name, field_value in custom_fields.items():
        if field_name in standard_fields:
            continue
        
        # フィールド名を見出しに
        doc.add_heading(field_name, level=1)
        
        if isinstance(field_value, list):
            # リスト型: 番号付きリスト
            for i, item in enumerate(field_value, 1):
                doc.add_paragraph(f'{i}. {item}', style='List Number')
        else:
            # 文字列型: そのまま表示
            doc.add_paragraph(str(field_value))
        
        doc.add_paragraph('')  # 空行
    
    # BytesIOに保存
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _generate_xlsx(data: DownloadRequest) -> BytesIO:
    """Generate an Excel document from minute data (mojiokoshi7.py準拠)"""
    from datetime import datetime
    
    buffer = BytesIO()
    
    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
        # シート1: メタデータ
        metadata = [
            {'項目': 'タイトル', '値': data.title},
            {'項目': '作成日時', '値': datetime.now().strftime('%Y年%m月%d日 %H:%M')},
            {'項目': '総セグメント数', '値': f'{len(data.segments)}件'},
            {'項目': '決定事項数', '値': f'{len(data.decisions)}件'},
            {'項目': 'アクションアイテム数', '値': f'{len(data.action_items)}件'},
        ]
        
        # 話者数を追加
        speakers = set(seg.speaker for seg in data.segments if seg.speaker)
        if speakers:
            metadata.append({'項目': '話者数', '値': f'{len(speakers)}名'})
            metadata.append({'項目': '話者一覧', '値': ', '.join(sorted(speakers))})
        
        # 総時間を追加
        if data.segments:
            total_duration = max(seg.end for seg in data.segments)
            minutes = int(total_duration // 60)
            seconds = int(total_duration % 60)
            metadata.append({'項目': '総時間', '値': f'{minutes}分{seconds}秒'})
        
        metadata_df = pd.DataFrame(metadata)
        metadata_df.to_excel(writer, index=False, sheet_name='メタデータ')
        
        # シート2: 要約
        summary_data = []
        if data.summary:
            # 要約を複数行に分割
            summary_lines = data.summary.split('\n')
            for line in summary_lines:
                if line.strip():
                    summary_data.append({'カテゴリ': '概要', '内容': line.strip()})
        
        if summary_data:
            summary_df = pd.DataFrame(summary_data)
            summary_df.to_excel(writer, index=False, sheet_name='概要')
        
        # シート3: 決定事項
        if data.decisions:
            decisions_data = [
                {'No': i+1, '決定事項': decision}
                for i, decision in enumerate(data.decisions)
            ]
            decisions_df = pd.DataFrame(decisions_data)
            decisions_df.to_excel(writer, index=False, sheet_name='決定事項')
        
        # シート4: アクションアイテム
        if data.action_items:
            actions_data = [
                {'No': i+1, 'アクションアイテム': action}
                for i, action in enumerate(data.action_items)
            ]
            actions_df = pd.DataFrame(actions_data)
            actions_df.to_excel(writer, index=False, sheet_name='アクションアイテム')
        
        # シート5: 文字起こし詳細（mojiokoshi7.py形式）
        if data.segments:
            transcript_data = []
            for i, seg in enumerate(data.segments, 1):
                start_time = f"{int(seg.start//60):02d}:{int(seg.start%60):02d}"
                end_time = f"{int(seg.end//60):02d}:{int(seg.end%60):02d}"
                transcript_data.append({
                    'No': i,
                    '開始時刻': start_time,
                    '終了時刻': end_time,
                    '継続時間(秒)': round(seg.end - seg.start, 2),
                    '話者': seg.speaker if seg.speaker else "不明",
                    '内容': seg.text
                })
            
            transcript_df = pd.DataFrame(transcript_data)
            transcript_df.to_excel(writer, index=False, sheet_name='文字起こし')
        
        # カスタムフィールド(動的にシートを追加)
        custom_fields = {}
        if hasattr(data, '__pydantic_extra__'):
            custom_fields = data.__pydantic_extra__ or {}
        
        standard_fields = {'title', 'summary', 'decisions', 'action_items', 'segments'}
        for field_name, field_value in custom_fields.items():
            if field_name in standard_fields:
                continue
            
            # シート名(Excelの制限で最大31文字)
            sheet_name = field_name[:31]
            
            if isinstance(field_value, list):
                # リスト型: Noと内容の2カラムテーブル
                custom_df = pd.DataFrame({
                    'No': range(1, len(field_value) + 1),
                    field_name: field_value
                })
                custom_df.to_excel(writer, index=False, sheet_name=sheet_name)
            else:
                # 文字列型: キーと値の2カラム
                custom_df = pd.DataFrame({
                    '項目': [field_name],
                    '値': [str(field_value)]
                })
                custom_df.to_excel(writer, index=False, sheet_name=sheet_name)
    
    buffer.seek(0)
    return buffer


@app.post("/api/download_minutes", summary="Download minutes as Word or Excel file")
async def download_minutes(request: DownloadRequest, format: str = "docx"):
    """
    Download minutes in Word (.docx) or Excel (.xlsx) format
    
    Args:
        request: DownloadRequest containing minute data
        format: File format - 'docx' or 'xlsx'
    """
    try:
        if format == "docx":
            buffer = _generate_docx(request)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"{request.title}.docx"
        elif format == "xlsx":
            buffer = _generate_xlsx(request)
            media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            filename = f"{request.title}.xlsx"
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported format: {format}")
        
        # URL encode the filename to support Japanese characters
        from urllib.parse import quote
        encoded_filename = quote(filename)
        
        return StreamingResponse(
            buffer,
            media_type=media_type,
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
            }
        )
    
    except Exception as e:
        logger.error(f"File generation failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to generate file: {str(e)}")

